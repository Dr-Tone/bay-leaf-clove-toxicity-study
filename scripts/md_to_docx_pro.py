import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name="Times New Roman", size=12):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    r.rPr.rFonts.set(qn('w:ascii'), font_name)

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    # This gets access to the document.xml.rels file and adds a new relation id
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color element
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Add underline element
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    # Join them all together and append to the paragraph
    new_run.append(rPr)
    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def clean_html(text):
    # Strip HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('&nbsp;', ' ')
    # Strip non-link brackets like [Online] or [PMC12345] if they aren't part of a link
    text = re.sub(r'(?<!\])\[([^\]]+)\](?!\()', r'\1', text)
    return text

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


from docx.oxml import parse_xml

def add_math_equation(doc, eq_str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omml = ""
    if "Z^2" in eq_str:
        omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t xml:space="preserve">n = </m:t></m:r><m:f><m:num><m:r><m:t>Z²p(1 - p)</m:t></m:r></m:num><m:den><m:r><m:t>d²</m:t></m:r></m:den></m:f></m:oMath></m:oMathPara>'''
    elif "1.96" in eq_str:
        omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t xml:space="preserve">n = </m:t></m:r><m:f><m:num><m:r><m:t>1.96² × 0.5 × (1 - 0.5)</m:t></m:r></m:num><m:den><m:r><m:t>0.05²</m:t></m:r></m:den></m:f><m:r><m:t xml:space="preserve"> = 384.16 ≈ 384</m:t></m:r></m:oMath></m:oMathPara>'''
    elif "Average" in eq_str:
        omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t xml:space="preserve">Average Drugs per Encounter = </m:t></m:r><m:f><m:num><m:r><m:t>Σ Total Drugs</m:t></m:r></m:num><m:den><m:r><m:t>Σ Total Encounters</m:t></m:r></m:den></m:f></m:oMath></m:oMathPara>'''
    elif "Generic" in eq_str:
        omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t xml:space="preserve">% Generic = (</m:t></m:r><m:f><m:num><m:r><m:t>No. of Generics</m:t></m:r></m:num><m:den><m:r><m:t>Total Drugs</m:t></m:r></m:den></m:f><m:r><m:t>) × 100</m:t></m:r></m:oMath></m:oMathPara>'''
    elif "Antibiotic" in eq_str:
        omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t xml:space="preserve">% Antibiotics = (</m:t></m:r><m:f><m:num><m:r><m:t>Encounters with Antibiotics</m:t></m:r></m:num><m:den><m:r><m:t>Total Encounters</m:t></m:r></m:den></m:f><m:r><m:t>) × 100</m:t></m:r></m:oMath></m:oMathPara>'''
    elif "Injection" in eq_str:
        omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t xml:space="preserve">% Injections = (</m:t></m:r><m:f><m:num><m:r><m:t>Encounters with Injections</m:t></m:r></m:num><m:den><m:r><m:t>Total Encounters</m:t></m:r></m:den></m:f><m:r><m:t>) × 100</m:t></m:r></m:oMath></m:oMathPara>'''
    elif "EML" in eq_str:
        omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t xml:space="preserve">% EML Adherence = (</m:t></m:r><m:f><m:num><m:r><m:t>Drugs in EML</m:t></m:r></m:num><m:den><m:r><m:t>Total Drugs</m:t></m:r></m:den></m:f><m:r><m:t>) × 100</m:t></m:r></m:oMath></m:oMathPara>'''
    elif "n_{final}" in eq_str:
        # Special case for the oversampling formula with proper subscript
        if "384" in eq_str:
            omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:sSub><m:e><m:r><m:t>n</m:t></m:r></m:e><m:sub><m:r><m:t>final</m:t></m:r></m:sub></m:sSub><m:r><m:t xml:space="preserve"> = 384 + (384 × 0.10) = 422</m:t></m:r></m:oMath></m:oMathPara>'''
        elif "422.4" in eq_str:
            omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:sSub><m:e><m:r><m:t>n</m:t></m:r></m:e><m:sub><m:r><m:t>final</m:t></m:r></m:sub></m:sSub><m:r><m:t xml:space="preserve"> = 422.4 ≈ 422</m:t></m:r></m:oMath></m:oMathPara>'''
        else:
            omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:sSub><m:e><m:r><m:t>n</m:t></m:r></m:e><m:sub><m:r><m:t>final</m:t></m:r></m:sub></m:sSub><m:r><m:t xml:space="preserve"> = n + (n × 0.10)</m:t></m:r></m:oMath></m:oMathPara>'''
    elif "r =" in eq_str or "Pearson" in eq_str:
        omml = '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t>r</m:t></m:r></m:oMath></m:oMathPara>'''
    else:
        return
        
    if omml:
        element = parse_xml(omml)
        p._p.append(element)

def build_docx(md_path, docx_path, image_path):
    doc = Document()
    
    # Global Style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Setup Heading Styles for TOC
    for h_level, size in [('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 12), ('Heading 4', 12)]:
        if h_level not in doc.styles:
            doc.styles.add_style(h_level, 1) # 1 is for WD_STYLE_TYPE.PARAGRAPH
        h_style = doc.styles[h_level]
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        h_font.size = Pt(size)
        h_font.bold = True
        h_font.italic = False
        h_font.color.rgb = RGBColor(0, 0, 0)
    
    # Setup initial section (Pages 1 & 2: Cover Pages)
    section = doc.sections[0]
    
    # Set Page Size to A4
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    
    # We leave the first section's footer empty so Pages 1 & 2 have no numbers.
    section.different_first_page_header_footer = False 
    
    # Ensure no footer paragraphs exist or are empty
    footer = section.footer
    for p in footer.paragraphs:
        p.text = ""
    
    with open(md_path, 'r', encoding='utf-8-sig') as f:
        lines = f.readlines()

    current_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    title_text = "EVALUATION OF ACUTE ORAL TOXICITY AND ANTIDYSMENORRHEA EFFECTS OF CLOVE (SYZYGIUM AROMATICUM)"
    title_seen = False
    
    i = 0
    while i < len(lines):
        line = lines[i].replace('﻿', '').strip().replace('\ufeff', '')
        
        if '<div' in line:
            # Handle Page Breaks
            if 'page-break-after: always' in line:
                doc.add_page_break()
            
            # Handle Alignment
            align_match = re.search(r'text-align:\s*(center|justify|left|right)', line)
            if align_match:
                align_val = align_match.group(1).lower()
                if align_val == 'center': current_alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align_val == 'justify': current_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else: current_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            i += 1
            continue
        elif '</div>' in line:
            current_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            i += 1
            continue
        elif '---' == line:
            doc.add_page_break()
            i += 1
            continue
        elif '<br' in line: 
            i += 1
            continue

        if not line:
            i += 1
            continue


        # Math Equations
        if line.startswith('$$') or line.startswith('$n_{final}'):
            add_math_equation(doc, line)
            i += 1
            continue

        # Tables
        if line.startswith('|') and i+1 < len(lines) and re.match(r'^\|[\s:\-\|]+$', lines[i+1].strip()):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for tline in table_lines:
                if re.match(r'^\|[\s:\-\|]+$', tline):
                    continue
                cells = [clean_html(c).strip() for c in tline.split('|')]
                if cells and not cells[0]: cells = cells[1:]
                if cells and not cells[-1]: cells = cells[:-1]
                rows.append(cells)
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                
                is_signature = any('___' in c for r in rows for c in r)
                
                if is_signature:
                    # No borders for signature tables
                    try:
                        table.style = 'Table Normal'
                    except:
                        pass
                else:
                    # Implement "Three-Line" (APA/Scholarly) Table Style
                    tbl = table._element
                    tblPr = tbl.xpath('w:tblPr')[0]
                    
                    # 1. Clear existing borders if any
                    existing_borders = tblPr.xpath('w:tblBorders')
                    for b in existing_borders:
                        tblPr.remove(b)
                        
                    # 2. Add Top, Bottom, and Inside Horizontal (below header) borders
                    tblBorders = OxmlElement('w:tblBorders')
                    
                    for border_name in ['top', 'bottom']:
                        b = OxmlElement(f'w:{border_name}')
                        b.set(qn('w:val'), 'single')
                        b.set(qn('w:sz'), '12') # 1.5 pt
                        b.set(qn('w:space'), '0')
                        b.set(qn('w:color'), '000000')
                        tblBorders.append(b)
                    
                    # Explicitly set other borders to 'nil' (transparent)
                    for border_name in ['left', 'right', 'insideV', 'insideH']:
                        b = OxmlElement(f'w:{border_name}')
                        b.set(qn('w:val'), 'nil')
                        tblBorders.append(b)
                        
                    tblPr.append(tblBorders)
                    
                for r_idx, row_data in enumerate(rows):
                    tr = table.rows[r_idx]._tr
                    if not is_signature and r_idx == 0:
                        # Add border below header row
                        trPr = tr.get_or_add_trPr()
                        # We apply it to the cells in the first row
                        for cell_idx in range(len(row_data)):
                            tc = table.cell(r_idx, cell_idx)._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcBorders = OxmlElement('w:tcBorders')
                            bottom = OxmlElement('w:bottom')
                            bottom.set(qn('w:val'), 'single')
                            bottom.set(qn('w:sz'), '6') # 0.75 pt
                            bottom.set(qn('w:color'), '000000')
                            tcBorders.append(bottom)
                            tcPr.append(tcBorders)

                    for c_idx, cell_data in enumerate(row_data):
                        if c_idx < len(table.columns):
                            cell = table.cell(r_idx, c_idx)
                            p = cell.paragraphs[0]
                            
                            if is_signature:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                            # Handle bolding and font
                            display_text = cell_data.replace('**', '')
                            run = p.add_run(display_text)
                            if '**' in cell_data or r_idx == 0: run.bold = True
                            set_font(run, size=12 if is_signature else 11)
            continue

        # Headers
        if line.startswith('#'):
            level = line.count('#')
            text = clean_html(line.replace('#', '')).strip()
            if not text:
                i += 1
                continue
                
            # Sentence Case logic + Page Breaks
            if level == 1:
                text = text.upper()
                
                # Handle duplicate title on second cover page
                if text == title_text:
                    if title_seen:
                        # Second occurrence: render as bold centered text, NOT as heading (so it won't appear in TOC)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_after = Pt(40)
                        run = p.add_run(text)
                        run.bold = True
                        set_font(run, size=14)
                        i += 1
                        continue
                    else:
                        title_seen = True
                
                # Page Break logic for Level 1
                no_break_headings = [
                    "INTRODUCTION", 
                    "LITERATURE REVIEW", 
                    "METHODOLOGY", 
                    "RESULTS AND DATA ANALYSIS", 
                    "DISCUSSION, CONCLUSION, AND RECOMMENDATIONS",
                ]
                
                if text == "CERTIFICATION":
                    # New section for Certification - start Roman numeral numbering here
                    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                    new_section.page_width = Inches(8.27)
                    new_section.page_height = Inches(11.69)
                    new_section.top_margin, new_section.bottom_margin = Inches(0.8), Inches(0.8)
                    new_section.left_margin, new_section.right_margin = Inches(0.85), Inches(0.85)
                    new_section.different_first_page_header_footer = False
                    new_section.footer.is_linked_to_previous = False
                    sectPr_cert = new_section._sectPr
                    pgNumType_cert = OxmlElement('w:pgNumType')
                    pgNumType_cert.set(qn('w:fmt'), 'lowerRoman')
                    pgNumType_cert.set(qn('w:start'), '1')
                    sectPr_cert.append(pgNumType_cert)
                    footer_cert = new_section.footer
                    p_footer_cert = footer_cert.paragraphs[0]
                    p_footer_cert.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_page_number(p_footer_cert.add_run())
                elif text == "CHAPTER ONE":
                    if len(doc.paragraphs) > 0:
                        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                        new_section.page_width = Inches(8.27)
                        new_section.page_height = Inches(11.69)
                        new_section.top_margin, new_section.bottom_margin, new_section.right_margin = Inches(0.8), Inches(0.8), Inches(0.85)
                        new_section.left_margin = Inches(0.85)
                        new_section.footer.is_linked_to_previous = False
                        sectPr2 = new_section._sectPr
                        pgNumType2 = OxmlElement('w:pgNumType')
                        pgNumType2.set(qn('w:fmt'), 'decimal')
                        pgNumType2.set(qn('w:start'), '1')
                        sectPr2.append(pgNumType2)
                        footer2 = new_section.footer
                        p_footer2 = footer2.paragraphs[0]
                        p_footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        add_page_number(p_footer2.add_run())
                elif text.startswith("CHAPTER "):
                    doc.add_page_break()
                elif text not in no_break_headings and len(doc.paragraphs) > 0:
                    doc.add_page_break()
                    
                p = doc.add_paragraph(style='Heading 1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if text == title_text:
                    p.paragraph_format.space_after = Pt(20)
            elif level >= 2:
                # Level 2+: Trust the casing in the Markdown file
                p_style = 'Heading 2' if level == 2 else ('Heading 3' if level == 3 else 'Heading 4')
                p = doc.add_paragraph(style=p_style)
                p.alignment = current_alignment
                
            # Native TOC injection
            if text == "TABLE OF CONTENTS":
                run = p.add_run(text)
                set_font(run, size=14)
                
                p_toc = doc.add_paragraph()
                run_toc = p_toc.add_run()
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                
                for elem in [fldChar1, instrText, fldChar2, fldChar3]: run_toc._r.append(elem)
                
                i += 1
                continue

            run = p.add_run(text)
            set_font(run, size=14 if level == 1 else 12)
            i += 1
            continue

        # Images (General Markdown support)
        img_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
        if img_match:
            caption = img_match.group(1)
            raw_path = img_match.group(2).replace('file://', '')
            
            # Special case for Logo (image1.png) which uses the passed image_path
            if 'media/image1.png' in raw_path:
                target_img = image_path
            else:
                target_img = raw_path
                if not os.path.isabs(target_img):
                    target_img = os.path.join(os.path.dirname(md_path), target_img)

            if os.path.exists(target_img):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(target_img, width=Inches(4.5))
                p_img.paragraph_format.space_before = Pt(12)
                
                if caption:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_text = caption
                    if not cap_text.lower().startswith("figure"):
                        cap_text = f"Figure: {cap_text}"
                    run_cap = p_cap.add_run(cap_text)
                    run_cap.italic = True
                    set_font(run_cap, size=11)
                    p_cap.paragraph_format.space_after = Pt(12)
            else:
                print(f"Warning: Image not found at {target_img}")
            i += 1
            continue

        # Bullets
        if line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = current_alignment
            p.paragraph_format.line_spacing = 1.5
            text = clean_html(line[2:]).strip()
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\[[^\]]+\]\([^\)]+\))', text)
            for part in parts:
                if not part: continue
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(clean_html(part[2:-2]))
                    run.bold = True
                    set_font(run)
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(clean_html(part[1:-1]))
                    run.italic = True
                    set_font(run)
                elif part.startswith('[') and '](' in part:
                    # Handle Markdown link: [text](url)
                    match = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', part)
                    if match:
                        link_text = match.group(1)
                        link_url = match.group(2)
                        add_hyperlink(p, link_url, link_text)
                else:
                    run = p.add_run(clean_html(part))
                    set_font(run)
            i += 1
            continue

        # Normal Paragraph
        raw_line = lines[i]
        text_content = raw_line.replace('﻿', '').rstrip('\n').replace('\ufeff', '')
        
        if not clean_html(text_content).strip(): 
            i += 1
            continue
            
        p = doc.add_paragraph()
        p.alignment = current_alignment
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        
        # Robust Spacing for Cover Pages
        clean_text = clean_html(text_content).strip()
        is_cover_section = (not title_seen or (title_seen and doc.sections[-1].footer.paragraphs[0].text == ""))
        
        if is_cover_section:
            # We are in the front cover or inner cover (Section 0)
            # Page 2 (Inner Cover) has NO logo, so it needs more spacing to look "full"
            has_logo = (doc.paragraphs[-2].text == "" if len(doc.paragraphs) > 2 else False) # Rough check if previous was logo
            
            # Since we can't easily know the page number in the script, 
            # we distinguish by content: Inner cover has "SUPERVISED BY" and "SUBMITTED TO"
            is_inner_cover = ("SUPERVISED BY" in clean_text or "SUBMITTED TO" in clean_text or "PARTIAL FULFILMENT" in clean_text or (title_seen and "JUNE, 2026" in clean_text))
            
            mult = 2.0 if is_inner_cover else 1.0 # Double spacing for inner cover to fill the page
            
            if clean_text == "**BY**":
                p.paragraph_format.space_before = Pt(15 * mult)
                p.paragraph_format.space_after = Pt(15 * mult)
            elif "IYARE" in clean_text:
                p.paragraph_format.space_after = Pt(5 * mult)
            elif "MAT. NUMBER" in clean_text:
                p.paragraph_format.space_after = Pt(10 * mult)
            elif "SUPERVISED BY" in clean_text:
                p.paragraph_format.space_before = Pt(15 * mult)
                p.paragraph_format.space_after = Pt(20 * mult)
            elif "A PROJECT SUBMITTED TO" in clean_text:
                p.paragraph_format.space_before = Pt(30 * mult)
            elif "IN PARTIAL FULFILMENT" in clean_text:
                p.paragraph_format.space_before = Pt(15 * mult)
            elif "JUNE, 2026" in clean_text:
                p.paragraph_format.space_before = Pt(40 * mult)

        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\[[^\]]+\]\([^\)]+\))', text_content)
        for part in parts:
            if not part: continue
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(clean_html(part[2:-2]))
                run.bold = True
                set_font(run)
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(clean_html(part[1:-1]))
                run.italic = True
                set_font(run)
            elif part.startswith('[') and '](' in part:
                # Handle Markdown link: [text](url)
                match = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', part)
                if match:
                    link_text = match.group(1)
                    link_url = match.group(2)
                    add_hyperlink(p, link_url, link_text)
            else:
                run = p.add_run(clean_html(part))
                set_font(run)
        i += 1

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path} with PRO formatting.")

if __name__ == "__main__":
    import sys
    md = sys.argv[1] if len(sys.argv) > 1 else "output/IYARE-PROJECT-DRAFT.MD"
    docx = sys.argv[2] if len(sys.argv) > 2 else "output/IYARE-PROJECT-DRAFT.DOCX"
    img = "output/media/image1.png"
    build_docx(md, docx, img)
